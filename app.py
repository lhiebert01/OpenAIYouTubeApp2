# Part 1: Imports and Initial Setup
import streamlit as st
from openai import OpenAI
import os
import re
from pathlib import Path
import tiktoken
from youtube_transcript_api import YouTubeTranscriptApi
from googleapiclient.discovery import build
from dotenv import load_dotenv
import base64
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from typing import Dict, List, Optional, Tuple
import logging
import httpx
import warnings

# Configure logging to only show WARNING and above
# Configure logging
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# Specifically silence httpx logging
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("urllib3").setLevel(logging.WARNING)
logging.getLogger("google").setLevel(logging.WARNING)
logging.getLogger("google_auth_httplib2").setLevel(logging.WARNING)
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Constants
class Config:
    MODEL_NAME = "gpt-4-turbo-preview"
    MAX_CONTEXT_LENGTH = 128000
    MAX_OUTPUT_TOKENS = 4096
    PAGE_TITLE = "🎥 OpenAI YouTube Video Summary "
    PAGE_ICON = "🎥"
    
    ANALYSIS_TYPES = {
        "Summary & Key Points": """Please provide a comprehensive summary of the video content and list the key points in a well-organized table format. Include:
            1. Overall summary (2-3 paragraphs)
            2. Key points in a table with columns for Topic and Description""",
        "Title Suggestions": "Based on the video content, please suggest 3-5 alternative titles that accurately reflect the main themes and content. Explain why each title would be appropriate.",
        "Quotes with Timestamps": "Please extract 5-10 significant quotes from the video. For each quote, provide the context and explain its significance.",
        "Key Terms & Definitions": "Please identify and define key terms, concepts, and jargon used in the video. Present them in a clear, organized format.",
        "All Analysis": """Please provide a comprehensive analysis of the video including:
            1. Overall summary (2-3 paragraphs)
            2. Key points in a table format
            3. 3-5 alternative title suggestions with explanations
            4. 5-10 significant quotes with context
            5. Key terms and definitions"""
    }

# Initialize Streamlit state
if 'initialized' not in st.session_state:
    st.session_state.initialized = False
    
# Part 2: YouTubeAnalyzer Class (First Half)
class YouTubeAnalyzer:
    def __init__(self):
        if not st.session_state.initialized:
            self.setup_streamlit()
            st.session_state.initialized = True
        self.load_environment()
        self.initialize_openai()
        self.initialize_session_state()
        self.apply_custom_css()

    def setup_streamlit(self):
        """Initialize Streamlit configuration"""
        st.set_page_config(
            page_title=Config.PAGE_TITLE,
            page_icon=Config.PAGE_ICON,
            layout="wide",
            initial_sidebar_state="expanded"
        )

    def load_environment(self):
        """Load environment variables"""
        load_dotenv(override=True)
        self.openai_api_key = os.getenv("OPENAI_API")
        self.google_api_key = os.getenv("YOUTUBE_API_KEY")
        
        if not self.openai_api_key or not self.google_api_key:
            missing_keys = []
            if not self.openai_api_key:
                missing_keys.append("OPENAI_API")
            if not self.google_api_key:
                missing_keys.append("YOUTUBE_API_KEY")
            raise ValueError(f"Missing required API keys: {', '.join(missing_keys)}")

    def initialize_openai(self):
        """Initialize OpenAI client"""
        self.client = OpenAI(api_key=self.openai_api_key)

    def initialize_session_state(self):
        """Initialize Streamlit session state variables"""
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
        if 'current_analysis' not in st.session_state:
            st.session_state.current_analysis = None
        if 'transcript' not in st.session_state:
            st.session_state.transcript = None
        if 'full_text' not in st.session_state:
            st.session_state.full_text = None
        if 'title' not in st.session_state:
            st.session_state.title = None

    def apply_custom_css(self):
        """Apply custom CSS styling"""
        st.markdown("""
            <style>
                .big-font { font-size:24px !important; font-weight: bold; }
                .medium-font { font-size:20px !important; font-weight: bold; }
                .stButton>button {
                    background-color: #2196F3;
                    color: white;
                    border-radius: 10px;
                    padding: 0.5rem 1rem;
                    font-weight: bold;
                    border: none;
                    transition: all 0.3s ease;
                }
                .stButton>button:hover {
                    background-color: #1976D2;
                    color: white;
                }
                .stButton>button:active {
                    background-color: #1565C0;
                    color: white !important;
                }
                .download-button {
                    background-color: #2196F3 !important;
                    color: white !important;
                }
                .sidebar-content {
                    padding: 1rem;
                    background-color: #f0f2f6;
                    border-radius: 10px;
                }
                .analysis-section {
                    padding: 1.5rem;
                    background-color: white;
                    border-radius: 10px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                .chat-message {
                    padding: 1rem;
                    margin: 0.5rem 0;
                    border-radius: 5px;
                }
                .user-message {
                    background-color: #f0f2f6;
                }
                .assistant-message {
                    background-color: #e6f3ff;
                }
            </style>
        """, unsafe_allow_html=True)

    @staticmethod
    def count_tokens(text: str) -> int:
        """Count tokens in text using tiktoken"""
        try:
            encoding = tiktoken.encoding_for_model(Config.MODEL_NAME)
            return len(encoding.encode(text))
        except Exception as e:
            logger.error(f"Error counting tokens: {e}")
            return 0

    @staticmethod
    def truncate_text_to_token_limit(text: str, max_tokens: int = 100000) -> str:
        """Truncate text to fit within token limit while maintaining coherent sentences"""
        try:
            encoding = tiktoken.encoding_for_model(Config.MODEL_NAME)
            tokens = encoding.encode(text)
            
            if len(tokens) <= max_tokens:
                return text
                
            truncated_tokens = tokens[:max_tokens]
            truncated_text = encoding.decode(truncated_tokens)
            
            last_period = truncated_text.rfind('.')
            if last_period > 0:
                truncated_text = truncated_text[:last_period + 1]
            
            return truncated_text
        except Exception as e:
            logger.error(f"Error truncating text: {e}")
            return text[:max_tokens * 4]

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filenames for cross-platform compatibility"""
        return Path(filename).stem.replace(' ', '_')
    
# Part 3: YouTubeAnalyzer Class (Second Half) and Main Execution

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filenames for cross-platform compatibility"""
        return Path(filename).stem.replace(' ', '_')
    
    # Add this method here
    @staticmethod
    def fetch_transcript_with_timestamps(video_id: str) -> List[Dict]:
        """Fetch the YouTube video transcript with timestamps"""
        try:
            return YouTubeTranscriptApi.get_transcript(video_id)
        except Exception as e:
            logger.error(f"Error fetching transcript: {e}")
            raise ValueError(f"Error fetching transcript: {str(e)}")

  
  
    def fetch_video_title(self, video_id: str) -> str:
        """Fetch the YouTube video title"""
        try:
            # Disable OAuthlib's HTTPS verification when running locally
            os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '0'
            
            # Build the YouTube service without using cache
            youtube = build('youtube', 'v3', 
                          developerKey=self.google_api_key,
                          cache_discovery=False)
            
            request = youtube.videos().list(part="snippet", id=video_id)
            response = request.execute()
            
            if not response.get("items"):
                raise ValueError("Video not found")
                
            return response["items"][0]["snippet"]["title"]
        except Exception as e:
            logger.error(f"Error fetching video title: {e}")
            raise ValueError(f"Error fetching video title: {str(e)}")
      

    def get_openai_response(self, prompt: str, transcript: str) -> str:
        """Get response from OpenAI API"""
        try:
            system_message = "You are an assistant who provides accurate, well-organized summaries for video transcripts that outline the key points, ideas, and any quotes, concepts, insights gleaned from the video."
            
            system_tokens = self.count_tokens(system_message)
            prompt_tokens = self.count_tokens(prompt)
            max_transcript_tokens = Config.MAX_CONTEXT_LENGTH - system_tokens - prompt_tokens - Config.MAX_OUTPUT_TOKENS - 1000
            
            truncated_transcript = self.truncate_text_to_token_limit(transcript, max_transcript_tokens)
            
            if len(truncated_transcript) < len(transcript):
                st.sidebar.warning("Note: The transcript was truncated to fit within token limits.")
            
            response = self.client.chat.completions.create(
                model=Config.MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": f"{prompt}\n\nTranscript:\n{truncated_transcript}"}
                ],
                max_tokens=Config.MAX_OUTPUT_TOKENS,
                temperature=0.7,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"OpenAI API error: {e}")
            raise ValueError(f"OpenAI API error: {str(e)}")

    def export_to_word(self, content: str, filename: str = "analysis_export.docx") -> Optional[str]:
        """Convert content to Word document and create download link"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(st.session_state.title, 0)
            title.alignment = 1
            
            # Add subtitle
            subtitle = doc.add_heading('Analysis Results', level=1)
            subtitle.alignment = 1
            
            doc.add_paragraph()
            
            # Add the analysis content
            doc.add_heading('Analysis', level=2)
            paragraphs = st.session_state.current_analysis.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.style = 'Normal'
            
            # Add Q&A Section if chat history exists
            if st.session_state.chat_history:
                doc.add_heading('Questions & Answers', level=2)
                for i in range(0, len(st.session_state.chat_history), 2):
                    if i + 1 < len(st.session_state.chat_history):
                        q_para = doc.add_paragraph()
                        q_para.add_run('Question: ').bold = True
                        q_para.add_run(st.session_state.chat_history[i]["content"])
                        
                        a_para = doc.add_paragraph()
                        a_para.add_run('Answer: ').bold = True
                        a_para.add_run(st.session_state.chat_history[i + 1]["content"])
                        
                        doc.add_paragraph()
            
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            b64 = base64.b64encode(bio.getvalue()).decode()
            href = f'''
            <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" 
               download="{filename}"
               style="text-decoration: none;">
                <button class="download-button" style="
                    padding: 0.5rem 1rem;
                    border: none;
                    border-radius: 10px;
                    cursor: pointer;
                    font-weight: bold;
                    display: inline-flex;
                    align-items: center;
                    gap: 0.5rem;
                ">
                    📎 Download as Word Document
                </button>
            </a>
            '''
            return href
                    
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            st.error(f"Error creating Word document: {str(e)}")
            return None

    def get_full_analysis_text(self) -> str:
        """Combine analysis and chat history into a single text"""
        output = []
        
        output.append(f"Video Title: {st.session_state.title}\n")
        output.append("ANALYSIS RESULTS")
        output.append("=" * 80 + "\n")
        output.append(st.session_state.current_analysis)
        output.append("\n" + "=" * 80 + "\n")
        
        if st.session_state.chat_history:
            output.append("\nQUESTIONS & ANSWERS")
            output.append("=" * 80 + "\n")
            for i in range(0, len(st.session_state.chat_history), 2):
                if i + 1 < len(st.session_state.chat_history):
                    output.append(f"Question: {st.session_state.chat_history[i]['content']}")
                    output.append(f"Answer: {st.session_state.chat_history[i + 1]['content']}\n")
        
        return "\n".join(output)

    def extract_video_id(self, url: str) -> Optional[str]:
        """Extract video ID from YouTube URL"""
        patterns = [
            r"(?:v=|\/)([0-9A-Za-z_-]{11}).*",
            r"(?:embed\/)([0-9A-Za-z_-]{11})",
            r"(?:youtu\.be\/)([0-9A-Za-z_-]{11})"
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        return None

    def run_analysis(self, video_url: str, analysis_type: str):
        """Run the video analysis"""
        video_id = self.extract_video_id(video_url)
        if not video_id:
            st.sidebar.error("Invalid YouTube URL. Please enter a valid URL.")
            return

        try:
            if not st.session_state.transcript:
                transcript_data = self.fetch_transcript_with_timestamps(video_id)
                st.session_state.transcript = transcript_data
                st.session_state.full_text = " ".join([entry['text'] for entry in transcript_data])
                st.session_state.title = self.fetch_video_title(video_id)

            with st.spinner("Analyzing video content..."):
                prompt = Config.ANALYSIS_TYPES[analysis_type]
                st.session_state.current_analysis = self.get_openai_response(prompt, st.session_state.full_text)

        except Exception as e:
            st.sidebar.error(f"Error: {str(e)}")

    def render_sidebar(self):
        """Render the sidebar content"""
        with st.sidebar:
            st.markdown('<p class="medium-font">📊 Analysis Controls</p>', unsafe_allow_html=True)
            st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
            
            st.markdown("### Enter YouTube URL")
            video_url = st.text_input(
                label="YouTube Video URL",
                placeholder="Paste YouTube URL here...",
                label_visibility="collapsed"
            )
            
            st.markdown("### Select Analysis Type")
            analysis_type = st.radio(
                label="Select the type of analysis to perform",
                options=list(Config.ANALYSIS_TYPES.keys()),
                index=0,
                key="analysis_type_radio"
            )
            
            col1, col2 = st.columns([3, 1])
            with col1:
                if st.button("🚀 Run Analysis", key="run_analysis", type="primary", use_container_width=True) and video_url:
                    self.run_analysis(video_url, analysis_type)
            
            if 'current_analysis' in st.session_state and st.session_state.current_analysis:
                self.render_export_options()
            
            st.markdown('</div>', unsafe_allow_html=True)
                
    def render_export_options(self):
        """Render export options in sidebar"""
        st.markdown("### Export Options")
        
        full_analysis = self.get_full_analysis_text()
        
        st.download_button(
            label="📥 Download Analysis (TXT)",
            data=full_analysis,
            file_name=f"{self.sanitize_filename(st.session_state.title)}_analysis.txt",
            mime="text/plain",
            use_container_width=True
        )
        
        word_doc_link = self.export_to_word(
            full_analysis,
            f"{self.sanitize_filename(st.session_state.title)}_analysis.docx"
        )
        if word_doc_link:
            st.markdown(word_doc_link, unsafe_allow_html=True)
        
        if st.session_state.full_text:
            st.download_button(
                label="📄 Download Transcript",
                data=st.session_state.full_text,
                file_name=f"{self.sanitize_filename(st.session_state.title)}_transcript.txt",
                mime="text/plain",
                use_container_width=True
            )

    def render_main_content(self):
        """Render the main content area"""
        st.markdown('<p class="big-font">🎥 OpenAI YouTube Analysis App</p>', unsafe_allow_html=True)
        st.markdown("""
            <div style='background-color: #f0f2f6; padding: 1rem; border-radius: 10px; margin-bottom: 2rem;'>
                Transform YouTube content into comprehensive insights using OpenAI's powerful analysis.
            </div>
        """, unsafe_allow_html=True)

        if st.session_state.title:
            st.markdown(f'<p class="medium-font">📺 Analyzing: {st.session_state.title}</p>', unsafe_allow_html=True)

        if st.session_state.current_analysis:
            self.render_analysis_results()

        if st.session_state.full_text:
            self.render_transcript()
            self.render_chat_interface()

    def render_analysis_results(self):
        """Render the analysis results section"""
        st.markdown('<div class="analysis-section">', unsafe_allow_html=True)
        st.markdown('<p class="medium-font">Analysis Results</p>', unsafe_allow_html=True)
        st.markdown(st.session_state.current_analysis)
        st.markdown('</div>', unsafe_allow_html=True)

    def render_transcript(self):
        """Render the transcript section"""
        with st.expander("📝 Show Transcript"):
            st.text_area(
                label="Video transcript text",
                value=st.session_state.full_text,
                height=200,
                label_visibility="collapsed"
            )

    def render_chat_interface(self):
        """Render the chat interface section"""
        st.markdown('<div class="analysis-section">', unsafe_allow_html=True)
        st.markdown('<p class="medium-font">💬 Ask Questions About the Video</p>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([4, 1])
        with col1:
            user_question = st.text_input(
                label="Question about the video",
                placeholder="Type your question here...",
                label_visibility="collapsed"
            )
        with col2:
            ask_button = st.button("🤔 Ask", key="ask_question", use_container_width=True)
        
        if user_question and ask_button:
            self.handle_user_question(user_question)

        self.render_chat_history()
        st.markdown('</div>', unsafe_allow_html=True)

    def handle_user_question(self, question: str):
        """Handle user questions about the video"""
        prompt = f"Based on the video transcript, please answer this question: {question}"
        answer = self.get_openai_response(prompt, st.session_state.full_text)
        
        st.session_state.chat_history.append({"role": "user", "content": question})
        st.session_state.chat_history.append({"role": "assistant", "content": answer})

    def render_chat_history(self):
        """Render the chat history"""
        for i in range(len(st.session_state.chat_history)-1, -1, -2):
            if i >= 0:
                st.markdown(f"""
                    <div class="chat-message user-message">
                        <strong>Question:</strong><br>
                        {st.session_state.chat_history[i-1]["content"]}
                    </div>
                    <div class="chat-message assistant-message">
                        <strong>Answer:</strong><br>
                        {st.session_state.chat_history[i]["content"]}
                    </div>
                """, unsafe_allow_html=True)

    def render_footer(self):
        """Render the footer section"""
        st.markdown("""
            <div style='text-align: center; margin-top: 2rem; padding: 1rem; background-color: #f0f2f6; border-radius: 10px;'>
                <p>Developed by Lindsay Hiebert with OpenAI ❤️ using Streamlit</p>
            </div>
        """, unsafe_allow_html=True)


def main():     
    try:
        analyzer = YouTubeAnalyzer()
        analyzer.render_sidebar()
        analyzer.render_main_content()
        analyzer.render_footer()
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        logger.error(f"Application error: {e}", exc_info=True)

if __name__ == "__main__":
    main()